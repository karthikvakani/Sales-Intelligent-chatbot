"""
exports/docx_exporter.py
Generates a professionally formatted Word (.docx) report.
"""

from datetime import datetime
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config.settings import EXPORT_DIR
from utils.logger import logger


BRAND_DARK = RGBColor(0x1A, 0x2B, 0x4A)
BRAND_MID = RGBColor(0x25, 0x63, 0xEB)
BRAND_ACCENT = RGBColor(0xF5, 0x9E, 0x0B)


def _set_heading_colour(paragraph, colour: RGBColor):
    for run in paragraph.runs:
        run.font.color.rgb = colour


def _add_horizontal_rule(doc: DocxDocument):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2563EB")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)


def _render_value_docx(doc: DocxDocument, value, indent: int = 0):
    """Recursively render report values into the Word document."""
    if isinstance(value, str):
        p = doc.add_paragraph(value)
        p.paragraph_format.left_indent = Inches(indent * 0.25)
        p.paragraph_format.space_after = Pt(4)

    elif isinstance(value, list):
        for item in value:
            if isinstance(item, dict):
                for k, v in item.items():
                    p = doc.add_paragraph(style="List Bullet")
                    run = p.add_run(f"{k.replace('_', ' ').title()}: ")
                    run.bold = True
                    run.font.color.rgb = BRAND_DARK
                    _render_value_docx(doc, v, indent + 1)
            elif isinstance(item, str):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(item)
                p.paragraph_format.left_indent = Inches(indent * 0.25)

    elif isinstance(value, dict):
        for k, v in value.items():
            p = doc.add_paragraph()
            run = p.add_run(f"{k.replace('_', ' ').title()}: ")
            run.bold = True
            run.font.color.rgb = BRAND_DARK
            p.paragraph_format.left_indent = Inches(indent * 0.25)
            _render_value_docx(doc, v, indent + 1)


def export_docx(report: dict, company: str, country: str) -> str:
    """Generate a Word report and return the file path."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = "".join(c if c.isalnum() else "_" for c in company)
    filename = f"sales_report_{safe_name}_{timestamp}.docx"
    filepath = EXPORT_DIR / filename

    doc = DocxDocument()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # ── Cover ─────────────────────────────────────────────────────────────────
    title_p = doc.add_heading("Sales Intelligence Report", level=0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_heading_colour(title_p, BRAND_DARK)

    subtitle = doc.add_paragraph(f"{company}  |  {country}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = BRAND_MID
    subtitle.runs[0].bold = True

    date_p = doc.add_paragraph(
        f"Generated: {datetime.now().strftime('%B %d, %Y %H:%M UTC')}"
    )
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.runs[0].font.size = Pt(10)
    date_p.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # Confidence badge
    confidence = report.get("confidence", "unknown")
    conf_text = {
        "high": "✅ Confidence: High",
        "medium": "⚠️ Confidence: Medium",
        "low": "⚠️ Confidence: Low — Limited Data Available",
    }.get(confidence, f"Confidence: {confidence}")
    conf_p = doc.add_paragraph(conf_text)
    conf_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_p.runs[0].font.size = Pt(11)
    conf_p.runs[0].bold = True

    _add_horizontal_rule(doc)

    # ── Sections ──────────────────────────────────────────────────────────────
    sections = report.get("sections", {})
    section_order = [
        "Company Overview",
        "Industry Information",
        "Recent News & Events",
        "Potential Business Opportunities",
        "Pain Points Identified",
        "Suggested Sales Approach",
        "Key Contacts Guidance",
        "Competitive Context",
        "Data Gaps",
    ]

    for section_name in section_order:
        value = sections.get(section_name)
        if value is None:
            continue

        h = doc.add_heading(section_name, level=1)
        _set_heading_colour(h, BRAND_DARK)
        _add_horizontal_rule(doc)
        _render_value_docx(doc, value)

    # ── Sources ───────────────────────────────────────────────────────────────
    sources = report.get("sources", [])
    if sources:
        doc.add_page_break()
        h = doc.add_heading("Sources & References", level=1)
        _set_heading_colour(h, BRAND_DARK)
        _add_horizontal_rule(doc)
        for i, src in enumerate(sources, 1):
            p = doc.add_paragraph(f"{i}. {src}")
            p.style = doc.styles["Normal"]
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # ── Warnings ──────────────────────────────────────────────────────────────
    warnings = report.get("warnings", [])
    if warnings:
        h = doc.add_heading("⚠️ Caveats & Limitations", level=2)
        _set_heading_colour(h, BRAND_ACCENT)
        for w in warnings:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(w)

    doc.save(str(filepath))
    logger.info(f"DOCX exported: {filepath}")
    return str(filepath)
